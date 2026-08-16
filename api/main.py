from __future__ import annotations

import csv
import io
import hashlib
import json
import os
import secrets
import sqlite3
import subprocess
import uuid
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Literal

from fastapi import Body, FastAPI, Header, HTTPException, Query
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import Response
from openpyxl import Workbook, load_workbook
from openpyxl.chart import Reference, ScatterChart, Series
from openpyxl.styles import Alignment, Font, PatternFill
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.graphics.shapes import Circle, Drawing, Line, Rect, String
from reportlab.graphics.charts.piecharts import Pie
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from pydantic import BaseModel, ConfigDict, Field, model_validator


ROOT = Path(__file__).resolve().parents[1]
DATA_DIR = Path(os.getenv("TRANSISAFE_DATA_DIR", ROOT / "api" / "data"))
DB_PATH = DATA_DIR / "analyses.sqlite3"
APP_VERSION = "2.1.0"
DATASET_VERSION = "2.1"


class AnalysisRequest(BaseModel):
    model_config = ConfigDict(extra="forbid", allow_inf_nan=False)

    transistor_id: str = Field(pattern=r"^[A-Za-z0-9_.-]+$", min_length=1, max_length=63)
    vds_v: float = Field(ge=0)
    id_a: float = Field(ge=0)
    mode: Literal["LINEAR", "SWITCHING"]
    pulse_duration_s: float = Field(gt=0)
    frequency_hz: float = Field(ge=0)
    duty_cycle: float = Field(gt=0, le=1)
    temperature_reference: Literal["AMBIENT", "CASE"]
    temperature_c: float = Field(ge=-273.15)
    rth_cs_k_per_w: float = Field(ge=0)
    rth_sa_k_per_w: float = Field(ge=0)
    safety_factor: float = Field(ge=1)
    e_on_j: float = Field(ge=0)
    e_off_j: float = Field(ge=0)
    gate_drive_voltage_v: float = Field(ge=0)

    @model_validator(mode="after")
    def validate_switching_inputs(self) -> "AnalysisRequest":
        if self.mode == "SWITCHING":
            if self.frequency_hz <= 0:
                raise ValueError("frequency_hz must be greater than zero in switching mode")
            if self.e_on_j + self.e_off_j <= 0:
                raise ValueError("Eon or Eoff is required in switching mode")
            if self.gate_drive_voltage_v <= 0:
                raise ValueError("gate-drive voltage is required in switching mode")
        return self

    def to_engine_line(self) -> str:
        values = (
            self.transistor_id,
            self.vds_v,
            self.id_a,
            self.mode,
            self.pulse_duration_s,
            self.frequency_hz,
            self.duty_cycle,
            self.temperature_reference,
            self.temperature_c,
            self.rth_cs_k_per_w,
            self.rth_sa_k_per_w,
            self.safety_factor,
            self.e_on_j,
            self.e_off_j,
            self.gate_drive_voltage_v,
        )
        return ";".join(str(value) for value in values) + "\n"


class ReportConfiguration(BaseModel):
    model_config = ConfigDict(extra="forbid")
    report_type: Literal["analysis", "management"] = "analysis"
    sections: list[Literal["overview", "operating", "soa", "losses", "thermal", "margins", "traceability"]] = Field(
        default_factory=lambda: ["overview", "operating", "soa", "losses", "thermal", "margins", "traceability"], min_length=1
    )
    embed_charts: bool = True
    embed_3d: bool = True
    show_limits: bool = True


class SaveAnalysisRequest(BaseModel):
    model_config = ConfigDict(extra="forbid")
    name: str = Field(min_length=1, max_length=160)
    input: AnalysisRequest
    result: dict[str, Any]
    parent_id: str | None = None
    report_config: ReportConfiguration | None = None


class BatchRow(BaseModel):
    source_row: int = Field(ge=2)
    analysis_name: str = Field(default="", max_length=160)
    input: AnalysisRequest


class BatchAnalyzeRequest(BaseModel):
    rows: list[BatchRow] = Field(min_length=1, max_length=5000)


class AuthRequest(BaseModel):
    model_config = ConfigDict(extra="forbid")
    email: str = Field(pattern=r"^[^\s@]+@[^\s@]+\.[^\s@]+$", max_length=254)
    password: str = Field(min_length=8, max_length=200)
    name: str | None = Field(default=None, max_length=120)


def engine_path() -> Path:
    configured = os.getenv("TRANSISAFE_ENGINE")
    candidates = [
        Path(configured) if configured else None,
        ROOT / "build" / "Release" / "transisafe_json.exe",
        ROOT / "build" / "transisafe_json.exe",
        ROOT / "build" / "transisafe_json",
    ]
    for candidate in candidates:
        if candidate and candidate.is_file():
            return candidate
    raise HTTPException(
        status_code=503,
        detail="The C analysis engine is not built. Build the CMake target transisafe_json first.",
    )


def load_models() -> list[dict[str, object]]:
    models: list[dict[str, object]] = []
    metadata: dict[str, dict[str, str]] = {}
    rds_on_25_c: dict[str, float] = {}
    with (ROOT / "component_metadata.csv").open(encoding="utf-8", newline="") as handle:
        for row in csv.DictReader(handle, delimiter=";"):
            metadata[row["transistor_id"]] = row
    with (ROOT / "mosfet_curves.csv").open(encoding="utf-8", newline="") as handle:
        rows = (line for line in handle if not line.lstrip().startswith("#"))
        for row in csv.DictReader(rows, delimiter=";"):
            if row["curve_type"] == "RDS_ON" and float(row["x"]) == 25:
                rds_on_25_c[row["transistor_id"]] = float(row["y"])
    with (ROOT / "transistors.csv").open(encoding="utf-8", newline="") as handle:
        rows = (line for line in handle if not line.lstrip().startswith("#"))
        for row in csv.DictReader(rows, delimiter=";"):
            presentation = metadata.get(row["transistor_id"], {})
            models.append(
                {
                    "id": row["transistor_id"],
                    "type": row["type"],
                    "vds_max_v": float(row["vds_max"]),
                    "id_continuous_max_a": float(row["id_continuous_max"]),
                    "id_pulse_max_a": float(row["id_pulse_max"]),
                    "tj_max_c": float(row["tj_max"]),
                    "rth_jc_k_per_w": float(row["rth_jc"]),
                    "rds_on_25_ohm": rds_on_25_c.get(row["transistor_id"]),
                    "datasheet_url": row["datasheet_url"],
                    "revision": row["datasheet_revision"],
                    "retrieved_date": row["retrieved_date"],
                    "manufacturer": presentation.get("manufacturer", ""),
                    "datasheet_type": presentation.get("datasheet_type", row["transistor_id"]),
                    "package_name": presentation.get("package_name", ""),
                    "product_url": presentation.get("product_url", ""),
                    "image_path": presentation.get("image_path", ""),
                    "development_fixture": row["transistor_id"] == "ENGINEERING_FIXTURE",
                    "verification_status": "REVIEW_PENDING",
                    "curve_status": "ENGINEERING_DIGITIZATION",
                }
            )
    return models


def load_soa_curves(transistor_id: str) -> list[dict[str, object]]:
    grouped: dict[float, list[dict[str, float]]] = {}
    with (ROOT / "mosfet_curves.csv").open(encoding="utf-8", newline="") as handle:
        rows = (line for line in handle if not line.lstrip().startswith("#"))
        for row in csv.DictReader(rows, delimiter=";"):
            if row["transistor_id"] != transistor_id or row["curve_type"] != "SOA":
                continue
            parameter = float(row["parameter"])
            grouped.setdefault(parameter, []).append(
                {"vds_v": float(row["x"]), "id_a": float(row["y"])}
            )
    return [
        {"pulse_duration_s": parameter, "points": sorted(points, key=lambda point: point["vds_v"])}
        for parameter, points in sorted(grouped.items())
    ]


def utc_now() -> str:
    return datetime.now(timezone.utc).replace(microsecond=0).isoformat()


def init_storage() -> None:
    DATA_DIR.mkdir(parents=True, exist_ok=True)
    with sqlite3.connect(DB_PATH) as connection:
        connection.execute(
            """CREATE TABLE IF NOT EXISTS analyses (
                id TEXT PRIMARY KEY,
                name TEXT NOT NULL,
                parent_id TEXT,
                input_json TEXT NOT NULL,
                result_json TEXT NOT NULL,
                created_at TEXT NOT NULL,
                updated_at TEXT NOT NULL
            )"""
        )
        connection.execute("""CREATE TABLE IF NOT EXISTS users (
            id TEXT PRIMARY KEY, email TEXT NOT NULL UNIQUE, name TEXT NOT NULL,
            password_hash TEXT NOT NULL, salt TEXT NOT NULL, created_at TEXT NOT NULL
        )""")
        connection.execute("""CREATE TABLE IF NOT EXISTS sessions (
            token TEXT PRIMARY KEY, user_id TEXT NOT NULL, created_at TEXT NOT NULL,
            FOREIGN KEY(user_id) REFERENCES users(id)
        )""")


def password_hash(password: str, salt: bytes) -> str:
    return hashlib.pbkdf2_hmac("sha256", password.encode(), salt, 210_000).hex()


def session_user(token: str) -> dict[str, str]:
    init_storage()
    with sqlite3.connect(DB_PATH) as connection:
        connection.row_factory=sqlite3.Row
        row=connection.execute("SELECT users.id, users.email, users.name FROM sessions JOIN users ON users.id=sessions.user_id WHERE sessions.token=?",(token,)).fetchone()
    if row is None: raise HTTPException(status_code=401,detail="Invalid or expired session")
    return dict(row)


def run_engine(request: AnalysisRequest) -> dict[str, Any]:
    try:
        completed = subprocess.run(
            [str(engine_path()), "--transistors", str(ROOT / "transistors.csv"),
             "--curves", str(ROOT / "mosfet_curves.csv")],
            input=request.to_engine_line(), text=True, capture_output=True,
            cwd=ROOT, timeout=10, check=False,
        )
    except (OSError, subprocess.TimeoutExpired) as exc:
        raise HTTPException(status_code=503, detail=f"Analysis engine unavailable: {exc}") from exc
    try:
        payload = json.loads(completed.stdout)
    except json.JSONDecodeError as exc:
        raise HTTPException(status_code=502, detail="The analysis engine returned an invalid response.") from exc
    if completed.returncode != 0 or not payload.get("ok"):
        message = payload.get("error", {}).get("message", "Analysis failed")
        raise HTTPException(status_code=422, detail=message)
    payload["analysis_metadata"] = {
        "timestamp": utc_now(), "engine_version": payload.get("source", {}).get("engine_version", APP_VERSION),
        "dataset_version": payload.get("source", {}).get("dataset_version", DATASET_VERSION),
        "application_version": APP_VERSION,
        "warnings": analysis_warnings(payload),
        "assumptions": ["Stored engineering curves are used unchanged by the native C engine."],
        "model_limitations": ["Engineering decision support; not certification or qualification evidence."],
    }
    return payload


def analysis_warnings(payload: dict[str, Any]) -> list[str]:
    result = payload.get("result", {})
    warnings: list[str] = []
    if not result.get("data_complete", False):
        warnings.append("Required engineering data is incomplete; no SAFE classification is permitted.")
    status = str(result.get("status", "UNKNOWN"))
    if status == "CRITICAL": warnings.append("At least one engineering reserve is below the configured critical threshold.")
    if status.startswith("NOT_SAFE"):
        warnings.append(f"Operating point violates the native engine constraint: {result.get('reason', 'UNKNOWN')}.")
    if payload.get("source", {}).get("verification_status") != "VERIFIED":
        warnings.append("Stored engineering digitization is pending independent review.")
    return warnings


def analysis_record(row: sqlite3.Row) -> dict[str, Any]:
    return {"id": row["id"], "name": row["name"], "parent_id": row["parent_id"],
            "input": json.loads(row["input_json"]), "result": json.loads(row["result_json"]),
            "created_at": row["created_at"], "updated_at": row["updated_at"]}


def get_analysis_or_404(analysis_id: str) -> dict[str, Any]:
    init_storage()
    with sqlite3.connect(DB_PATH) as connection:
        connection.row_factory = sqlite3.Row
        row = connection.execute("SELECT * FROM analyses WHERE id = ?", (analysis_id,)).fetchone()
    if row is None: raise HTTPException(status_code=404, detail="Saved analysis not found")
    return analysis_record(row)


def status_fill(status: str) -> PatternFill:
    color = "DDF4E8" if status in {"SAFE", "PASS"} else "FFF1CC" if status in {"CRITICAL", "WARNING"} else "FDE2E5"
    return PatternFill("solid", fgColor=color)


def add_sheet(workbook: Workbook, title: str, rows: list[list[Any]], widths: tuple[int, ...] = (30, 24, 24, 48)):
    sheet = workbook.create_sheet(title)
    for row in rows: sheet.append(row)
    sheet.freeze_panes = "A2"
    sheet.auto_filter.ref = sheet.dimensions
    for cell in sheet[1]:
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="0B5CAD")
        cell.alignment = Alignment(vertical="center")
    for index, width in enumerate(widths, 1): sheet.column_dimensions[chr(64 + index)].width = width
    return sheet


def workbook_for_analysis(record: dict[str, Any]) -> bytes:
    payload, inp = record["result"], record["input"]
    result, source = payload["result"], payload["source"]
    meta = payload.get("analysis_metadata", {})
    model = next((item for item in load_models() if item["id"] == inp["transistor_id"]), {})
    workbook = Workbook(); workbook.remove(workbook.active)
    summary = add_sheet(workbook, "SUMMARY", [["Field", "Value", "Unit"],
        ["Analysis ID", record.get("id", "")], ["Analysis Name", record.get("name", "")],
        ["Timestamp", meta.get("timestamp", record.get("created_at", ""))], ["Component / OPN", inp["transistor_id"]],
        ["Manufacturer", model.get("manufacturer", "")], ["Operating Mode", inp["mode"]],
        ["Assessment", result["status"]], ["Closest Constraint", result["closest_constraint"]["type"]],
        ["Closest Reserve", result["closest_constraint"]["reserve_percent"], "%"], ["Tj", result["tj_c"], "°C"],
        ["Thermal Reserve", result["temperature_margin_c"], "K"], ["Total Loss", result["p_total_w"], "W"],
        ["SOA Reserve", result["margins"]["soa_reserve_percent"], "%"],
        ["Voltage Reserve", result["margins"]["voltage_reserve_percent"], "%"],
        ["Current Reserve", result["margins"]["current_reserve_percent"], "%"],
        ["Warnings", " | ".join(meta.get("warnings", []))], ["Data Coverage", "Complete" if result["data_complete"] else "Incomplete"]])
    summary["B8"].fill = status_fill(str(result["status"]))
    add_sheet(workbook, "INPUTS", [["Input", "Value", "Unit"],
        ["VDS", inp["vds_v"], "V"], ["ID", inp["id_a"], "A"], ["Pulse Duration", inp["pulse_duration_s"], "s"],
        ["Duty Cycle", inp["duty_cycle"], "ratio"], ["Frequency", inp["frequency_hz"], "Hz"],
        ["Gate Voltage", inp["gate_drive_voltage_v"], "V"], ["Eon", inp["e_on_j"], "J"], ["Eoff", inp["e_off_j"], "J"],
        ["Temperature", inp["temperature_c"], "°C"], ["Thermal Reference", inp["temperature_reference"], ""],
        ["RθCS", inp["rth_cs_k_per_w"], "K/W"], ["RθSA", inp["rth_sa_k_per_w"], "K/W"], ["Safety Factor", inp["safety_factor"], "×"]])
    checks = result["checks"]
    add_sheet(workbook, "ENGINEERING CHECKS", [["Check", "Input / Assessment", "Engineering Basis", "Limit", "Margin", "Status"],
        ["Voltage", inp["vds_v"], "Native rating check incl. safety factor", source["vds_max_v"], result["margins"]["voltage_reserve_percent"], "PASS" if checks["voltage"] else "FAIL"],
        ["Current", inp["id_a"], "Native applicable current limit", result["current_limit_a"], result["margins"]["current_reserve_percent"], "PASS" if checks["current"] else "FAIL"],
        ["Safe Operating Area", inp["id_a"], "Stored SOA interpolation", result["soa_limit_a"], result["margins"]["soa_reserve_percent"], "PASS" if checks["soa"] else "FAIL"],
        ["Junction Temperature", result["tj_c"], "Native thermal path", source["tj_max_c"], result["temperature_margin_c"], "PASS" if checks["temperature"] else "FAIL"],
        ["Data Coverage", "Complete" if result["data_complete"] else "Incomplete", "Required native model and curves", "Complete", "", "PASS" if result["data_complete"] else "UNKNOWN"]], (24,24,42,18,18,16))
    add_sheet(workbook, "LIMITS", [["Condition", "Engine-provided Limit", "Unit", "Source"],
        [f"At VDS = {inp['vds_v']} V", payload["optimization"]["max_current_a"] if payload["optimization"]["max_current_available"] else None, "A", "Native iterative analysis"],
        [f"At ID = {inp['id_a']} A", payload["optimization"]["max_voltage_v"] if payload["optimization"]["max_voltage_available"] else None, "V", "Native iterative analysis"]])
    losses = [("Pcond", result["p_conduction_w"]), ("Psw", result["p_switching_w"]), ("Pgate", result["p_gate_w"])]
    total = result["p_total_w"]
    add_sheet(workbook, "LOSSES", [["Loss", "Value", "Unit", "Percentage of Total"]] + [[name, value, "W", value / total * 100 if total > 0 else None] for name, value in losses] + [["Ptotal", total, "W", 100 if total > 0 else None]])
    add_sheet(workbook, "THERMAL", [["Node / Property", "Value", "Unit"], ["Tj", result["tj_c"], "°C"],
        [inp["temperature_reference"], inp["temperature_c"], "°C"], ["ΔT", result["tj_c"] - inp["temperature_c"], "K"],
        ["Tj max", source["tj_max_c"], "°C"], ["Thermal reserve", result["temperature_margin_c"], "K"],
        ["ZthJC", result["zth_jc_k_per_w"], "K/W"], ["Thermal mode", inp["temperature_reference"], ""]])
    soa_curves=load_soa_curves(inp["transistor_id"])
    soa = add_sheet(workbook, "SOA", [["Pulse Duration [s]", "VDS [V]", "ID Boundary [A]"]] +
        [[curve["pulse_duration_s"], point["vds_v"], point["id_a"]] for curve in soa_curves for point in curve["points"]] +
        [["OPERATING_POINT", inp["vds_v"], inp["id_a"]]])
    soa.sheet_view.showGridLines = False
    chart=ScatterChart();chart.title="Stored Safe-Operating-Area Curves";chart.x_axis.title="VDS [V]";chart.y_axis.title="ID [A]";chart.x_axis.scaling.logBase=10;chart.y_axis.scaling.logBase=10;chart.height=9;chart.width=16
    cursor=2
    for curve in soa_curves:
        end=cursor+len(curve["points"])-1;xvalues=Reference(soa,min_col=2,min_row=cursor,max_row=end);yvalues=Reference(soa,min_col=3,min_row=cursor,max_row=end)
        series=Series(yvalues,xvalues,title=f"{curve['pulse_duration_s']:g} s");series.graphicalProperties.line.width=18000;chart.series.append(series);cursor=end+1
    point_series=Series(Reference(soa,min_col=3,min_row=cursor,max_row=cursor),Reference(soa,min_col=2,min_row=cursor,max_row=cursor),title="Operating Point");point_series.marker.symbol="circle";point_series.marker.size=9;point_series.graphicalProperties.line.noFill=True;chart.series.append(point_series);soa.add_chart(chart,"E2")
    add_sheet(workbook, "TRACEABILITY", [["Field", "Value"], ["Component / OPN", inp["transistor_id"]],
        ["Manufacturer", model.get("manufacturer", "")], ["Package", model.get("package_name", "")], ["Datasheet", source["datasheet_url"]],
        ["Datasheet Revision", source["revision"]], ["Retrieved Date", source["retrieved_date"]], ["Dataset Version", source["dataset_version"]],
        ["Curve Status", source["curve_status"]], ["Verification Status", source["verification_status"]], ["C-Engine Version", source["engine_version"]],
        ["Application Version", source["application_version"]]])
    add_sheet(workbook, "METADATA", [["Field", "Value"], ["Analysis ID", record.get("id", "")], ["Timestamp", meta.get("timestamp", "")],
        ["Engine Version", meta.get("engine_version", "")], ["Dataset Revision", meta.get("dataset_version", "")], ["Analysis Mode", inp["mode"]],
        ["Safety Factor", inp["safety_factor"]], ["Warnings", " | ".join(meta.get("warnings", []))],
        ["Assumptions", " | ".join(meta.get("assumptions", []))], ["Model Limitations", " | ".join(meta.get("model_limitations", []))]])
    config = record.get("report_config") or {}
    selected = set(config.get("sections") or ["overview", "operating", "soa", "losses", "thermal", "margins", "traceability"])
    sheet_sections = {"SUMMARY": "overview", "INPUTS": "operating", "ENGINEERING CHECKS": "margins", "LIMITS": "margins",
                      "LOSSES": "losses", "THERMAL": "thermal", "SOA": "soa", "TRACEABILITY": "traceability", "METADATA": "traceability"}
    for sheet_name, section_name in sheet_sections.items():
        if section_name not in selected and sheet_name in workbook.sheetnames:
            workbook.remove(workbook[sheet_name])
    stream = io.BytesIO(); workbook.save(stream); return stream.getvalue()


def pdf_for_analysis(record: dict[str, Any]) -> bytes:
    payload, inp = record["result"], record["input"]
    result, source = payload["result"], payload["source"]
    stream = io.BytesIO(); styles = getSampleStyleSheet()
    document = SimpleDocTemplate(stream, pagesize=A4, rightMargin=18*mm, leftMargin=18*mm, topMargin=16*mm, bottomMargin=16*mm,
                                 title=f"TransiSafe Engineering Report – {record.get('name', '')}")
    story = [Paragraph("TransiSafe", styles["Title"]), Paragraph("ENGINEERING DECISION SUPPORT", styles["Heading3"]), Spacer(1, 8),
             Paragraph(f"Engineering Report · {record.get('name', 'Analysis')}", styles["Heading1"])]
    rows = [["Analysis ID", record.get("id", "")], ["Component", inp["transistor_id"]], ["Mode", inp["mode"]],
            ["Operating Point", f"{inp['vds_v']} V · {inp['id_a']} A · {inp['temperature_c']} °C"],
            ["Assessment", result["status"]], ["Closest Constraint", f"{result['closest_constraint']['type']} · {result['closest_constraint']['reserve_percent']:.1f} % reserve"],
            ["Junction Temperature", f"{result['tj_c']:.1f} °C · {result['temperature_margin_c']:.1f} K reserve"],
            ["Total Loss", f"{result['p_total_w']:.3f} W"], ["SOA", f"{result['soa_limit_a']:.3f} A stored boundary"],
            ["Datasheet", source["revision"]], ["Dataset / Engine", f"{source['dataset_version']} / {source['engine_version']}"],
            ["Timestamp", payload.get("analysis_metadata", {}).get("timestamp", "")]]
    table = Table(rows, colWidths=[48*mm, 115*mm]); table.setStyle(TableStyle([("GRID",(0,0),(-1,-1),0.35,colors.HexColor("#CFDCEB")),
        ("BACKGROUND",(0,0),(0,-1),colors.HexColor("#EEF5FC")), ("FONTNAME",(0,0),(0,-1),"Helvetica-Bold"), ("VALIGN",(0,0),(-1,-1),"TOP"),
        ("LEFTPADDING",(0,0),(-1,-1),7), ("RIGHTPADDING",(0,0),(-1,-1),7), ("TOPPADDING",(0,0),(-1,-1),6), ("BOTTOMPADDING",(0,0),(-1,-1),6)])); story += [table, Spacer(1,12)]
    warnings = payload.get("analysis_metadata", {}).get("warnings", [])
    story += [Spacer(1,12),Paragraph("Engineering visual evidence",styles["Heading2"]),pdf_visuals(inp,result),Spacer(1,10)]
    story.append(Paragraph("Warnings and limitations", styles["Heading2"]))
    for warning in warnings: story.append(Paragraph(f"• {warning}", styles["BodyText"]))
    story.append(Paragraph("This report supports engineering decisions. It does not imply certification or component qualification.", styles["BodyText"]))
    document.build(story); return stream.getvalue()


def pdf_visuals(inp: dict[str,Any], result: dict[str,Any]) -> Drawing:
    drawing=Drawing(470,190);drawing.add(String(12,174,"SOA operating map",fontName="Helvetica-Bold",fontSize=10,fillColor=colors.HexColor("#173A5E")))
    curves=load_soa_curves(inp["transistor_id"]); points=[point for curve in curves for point in curve["points"]]
    if points:
        import math
        min_x,max_x=min(p["vds_v"] for p in points),max(p["vds_v"] for p in points);min_y,max_y=min(p["id_a"] for p in points),max(p["id_a"] for p in points)
        def xy(x:float,y:float)->tuple[float,float]: return (18+205*(math.log10(x)-math.log10(min_x))/max(.001,math.log10(max_x)-math.log10(min_x)),25+125*(math.log10(y)-math.log10(min_y))/max(.001,math.log10(max_y)-math.log10(min_y)))
        drawing.add(Rect(18,25,205,125,strokeColor=colors.HexColor("#D7E2EE"),fillColor=colors.white))
        palette=["#126BC5","#159CB1","#16945E","#D99A2B","#7758C7","#B94D62"]
        for index,curve in enumerate(curves):
            line_points=curve["points"]
            for a,b in zip(line_points,line_points[1:]):
                x1,y1=xy(a["vds_v"],a["id_a"]);x2,y2=xy(b["vds_v"],b["id_a"]);drawing.add(Line(x1,y1,x2,y2,strokeColor=colors.HexColor(palette[index%len(palette)]),strokeWidth=1.3))
        px,py=xy(max(min_x,min(max_x,inp["vds_v"])),max(min_y,min(max_y,inp["id_a"])));drawing.add(Circle(px,py,4,fillColor=colors.HexColor("#075BB8"),strokeColor=colors.white))
    drawing.add(String(255,174,"Power loss breakdown",fontName="Helvetica-Bold",fontSize=10,fillColor=colors.HexColor("#173A5E")))
    losses=[result["p_conduction_w"],result["p_switching_w"],result["p_gate_w"]];pie=Pie();pie.x=260;pie.y=42;pie.width=105;pie.height=105;pie.data=losses;pie.slices.strokeWidth=.5;pie.slices[0].fillColor=colors.HexColor("#126BC5");pie.slices[1].fillColor=colors.HexColor("#159CB1");pie.slices[2].fillColor=colors.HexColor("#7758C7");drawing.add(pie)
    total=result["p_total_w"]
    for i,(label,value,color) in enumerate(zip(["Conduction","Switching","Gate drive"],losses,["#126BC5","#159CB1","#7758C7"])):
        y=120-i*32;drawing.add(Rect(378,y,8,8,fillColor=colors.HexColor(color),strokeColor=None));drawing.add(String(391,y,label,fontSize=7));drawing.add(String(391,y-11,f"{value:.3f} W · {(value/total*100 if total else 0):.1f}%",fontName="Helvetica-Bold",fontSize=7))
    drawing.add(String(285,24,f"Total loss {total:.3f} W",fontName="Helvetica-Bold",fontSize=9,fillColor=colors.HexColor("#173A5E")))
    return drawing


app = FastAPI(
    title="TransiSafe Analysis API",
    version=APP_VERSION,
    description="HTTP adapter for the native TransiSafe MOSFET analysis core.",
)

allowed_origins = os.getenv(
    "TRANSISAFE_WEB_ORIGINS", "http://localhost:5173,http://localhost:4173"
).split(",")
app.add_middleware(
    CORSMiddleware,
    allow_origins=[origin.strip() for origin in allowed_origins if origin.strip()],
    allow_credentials=False,
    allow_methods=["GET", "POST", "PUT", "DELETE"],
    allow_headers=["Content-Type"],
)


@app.post("/api/auth/register", status_code=201)
def register(request: AuthRequest) -> dict[str, Any]:
    if not request.name or not request.name.strip(): raise HTTPException(status_code=422,detail="Name is required")
    init_storage(); user_id=str(uuid.uuid4()); salt=secrets.token_bytes(16); timestamp=utc_now()
    try:
        with sqlite3.connect(DB_PATH) as connection:
            connection.execute("INSERT INTO users VALUES (?,?,?,?,?,?)",(user_id,request.email.lower(),request.name.strip(),password_hash(request.password,salt),salt.hex(),timestamp))
    except sqlite3.IntegrityError as exc: raise HTTPException(status_code=409,detail="Account already exists") from exc
    return create_session(user_id,request.email.lower(),request.name.strip())


def create_session(user_id: str,email: str,name: str) -> dict[str,Any]:
    token=secrets.token_urlsafe(32)
    with sqlite3.connect(DB_PATH) as connection: connection.execute("INSERT INTO sessions VALUES (?,?,?)",(token,user_id,utc_now()))
    return {"token":token,"user":{"id":user_id,"email":email,"name":name}}


@app.post("/api/auth/login")
def login(request: AuthRequest) -> dict[str,Any]:
    init_storage()
    with sqlite3.connect(DB_PATH) as connection:
        connection.row_factory=sqlite3.Row; row=connection.execute("SELECT * FROM users WHERE email=?",(request.email.lower(),)).fetchone()
    if row is None or not secrets.compare_digest(row["password_hash"],password_hash(request.password,bytes.fromhex(row["salt"]))):
        raise HTTPException(status_code=401,detail="Invalid email or password")
    return create_session(row["id"],row["email"],row["name"])


@app.get("/api/auth/me")
def auth_me(authorization: str = Header(default="")) -> dict[str,Any]:
    if not authorization.startswith("Bearer "): raise HTTPException(status_code=401,detail="Missing session")
    return {"user":session_user(authorization[7:])}


@app.post("/api/auth/logout", status_code=204)
def logout(authorization: str = Header(default="")) -> Response:
    if authorization.startswith("Bearer "):
        init_storage()
        with sqlite3.connect(DB_PATH) as connection: connection.execute("DELETE FROM sessions WHERE token=?",(authorization[7:],))
    return Response(status_code=204)


@app.get("/api/health")
def health() -> dict[str, object]:
    available = True
    try:
        path = engine_path()
    except HTTPException:
        available = False
        path = None
    return {"ok": True, "engine_available": available, "engine": path.name if path else None,
            "engine_version": APP_VERSION, "dataset_version": DATASET_VERSION}


@app.get("/api/models")
def models() -> dict[str, object]:
    try:
        return {"models": load_models()}
    except (OSError, KeyError, ValueError) as exc:
        raise HTTPException(status_code=500, detail=f"Model database error: {exc}") from exc


@app.get("/api/models/{transistor_id}/soa-curves")
def soa_curves(transistor_id: str) -> dict[str, object]:
    if not transistor_id or len(transistor_id) > 63 or not all(
        character.isalnum() or character in "_.-" for character in transistor_id
    ):
        raise HTTPException(status_code=422, detail="Invalid transistor ID")
    model_ids = {model["id"] for model in load_models()}
    if transistor_id not in model_ids:
        raise HTTPException(status_code=404, detail="Unknown transistor model")
    try:
        return {
            "transistor_id": transistor_id,
            "curve_type": "SOA",
            "curves": load_soa_curves(transistor_id),
            "source": "Stored engineering digitization; no browser-side SOA calculation",
        }
    except (OSError, KeyError, ValueError) as exc:
        raise HTTPException(status_code=500, detail=f"Curve database error: {exc}") from exc


@app.post("/api/analyze")
def analyze(request: AnalysisRequest) -> dict[str, object]:
    return run_engine(request)


@app.get("/api/analyses")
def list_analyses() -> dict[str, Any]:
    init_storage()
    with sqlite3.connect(DB_PATH) as connection:
        connection.row_factory = sqlite3.Row
        rows = connection.execute("SELECT * FROM analyses ORDER BY updated_at DESC").fetchall()
    return {"analyses": [analysis_record(row) for row in rows]}


@app.get("/api/management/overview")
def management_overview() -> dict[str, Any]:
    records=list_analyses()["analyses"]; constraints:dict[str,int]={}; components:dict[str,int]={}; attention=0
    for record in records:
        result=record["result"].get("result",{});status=str(result.get("status","UNKNOWN"));constraint=str(result.get("closest_constraint",{}).get("type","UNKNOWN"));component=str(record["input"].get("transistor_id","UNKNOWN"))
        constraints[constraint]=constraints.get(constraint,0)+1;components[component]=components.get(component,0)+1
        if status != "SAFE": attention+=1
    return {"analysis_count":len(records),"attention_required":attention,"constraint_categories":constraints,"component_exposure":components,
            "source":"Saved native-engine analysis results; no management score is synthesized."}


@app.post("/api/analyses", status_code=201)
def save_analysis(request: SaveAnalysisRequest) -> dict[str, Any]:
    init_storage(); analysis_id = str(uuid.uuid4()); timestamp = utc_now()
    with sqlite3.connect(DB_PATH) as connection:
        connection.execute("INSERT INTO analyses VALUES (?, ?, ?, ?, ?, ?, ?)",
            (analysis_id, request.name, request.parent_id, request.input.model_dump_json(),
             json.dumps(request.result), timestamp, timestamp))
    return get_analysis_or_404(analysis_id)


@app.get("/api/analyses/{analysis_id}")
def get_analysis(analysis_id: str) -> dict[str, Any]: return get_analysis_or_404(analysis_id)


@app.put("/api/analyses/{analysis_id}")
def update_analysis(analysis_id: str, request: SaveAnalysisRequest) -> dict[str, Any]:
    get_analysis_or_404(analysis_id); timestamp = utc_now()
    with sqlite3.connect(DB_PATH) as connection:
        connection.execute("UPDATE analyses SET name=?, parent_id=?, input_json=?, result_json=?, updated_at=? WHERE id=?",
            (request.name, request.parent_id, request.input.model_dump_json(), json.dumps(request.result), timestamp, analysis_id))
    return get_analysis_or_404(analysis_id)


@app.delete("/api/analyses/{analysis_id}", status_code=204)
def delete_analysis(analysis_id: str) -> Response:
    get_analysis_or_404(analysis_id)
    with sqlite3.connect(DB_PATH) as connection: connection.execute("DELETE FROM analyses WHERE id=?", (analysis_id,))
    return Response(status_code=204)


@app.post("/api/analyses/{analysis_id}/duplicate", status_code=201)
def duplicate_analysis(analysis_id: str, name: str | None = Query(default=None, max_length=160)) -> dict[str, Any]:
    original = get_analysis_or_404(analysis_id)
    request = SaveAnalysisRequest(name=name or f"{original['name']} – Copy", input=AnalysisRequest.model_validate(original["input"]),
                                  result=original["result"], parent_id=analysis_id)
    return save_analysis(request)


@app.post("/api/analyses/{analysis_id}/rerun")
def rerun_analysis(analysis_id: str) -> dict[str, Any]:
    original = get_analysis_or_404(analysis_id); payload = run_engine(AnalysisRequest.model_validate(original["input"]))
    request = SaveAnalysisRequest(name=original["name"], input=AnalysisRequest.model_validate(original["input"]), result=payload,
                                  parent_id=original.get("parent_id"))
    return update_analysis(analysis_id, request)


def docx_for_analysis(record: dict[str, Any]) -> bytes:
    payload = record["result"]
    result = payload["result"]
    analysis_input = record["input"]
    document = Document()
    section = document.sections[0]
    section.top_margin = section.bottom_margin = Pt(42)
    section.left_margin = section.right_margin = Pt(48)

    title = document.add_heading("TransiSafe", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(10, 110, 209)
    subtitle = document.add_paragraph("ENGINEERING DECISION SUPPORT")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].bold = True
    subtitle.runs[0].font.size = Pt(9)
    document.add_heading(record.get("name", "Analysebericht"), level=1)
    document.add_paragraph(f"Generiert: {datetime.now(timezone.utc).astimezone().strftime('%d.%m.%Y %H:%M %Z')}")

    status = document.add_paragraph()
    run = status.add_run(f"Ergebnis: {result['status']}")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(15, 145, 85) if result["status"] == "SAFE" else RGBColor(197, 53, 53)
    status.add_run(f"\n{result.get('reason', '')}")

    document.add_heading("Ergebnisübersicht", level=2)
    summary = document.add_table(rows=1, cols=2)
    summary.style = "Light Shading Accent 1"
    summary.rows[0].cells[0].text = "Kennzahl"
    summary.rows[0].cells[1].text = "Wert"
    summary_rows = [
        ("Komponente", analysis_input["transistor_id"]),
        ("Betriebsart", analysis_input["mode"]),
        ("Betriebspunkt", f"{analysis_input['vds_v']:.3g} V / {analysis_input['id_a']:.3g} A"),
        ("Tj", f"{result['tj_c']:.1f} °C"),
        ("Gesamtverlust", f"{result['p_total_w']:.3f} W"),
        ("SOA-Reserve", f"{result['margins']['soa_reserve_percent']:.1f} %"),
        ("Thermische Reserve", f"{result['temperature_margin_c']:.1f} K"),
        ("Nächste Grenze", f"{result['closest_constraint']['type']} ({result['closest_constraint']['reserve_percent']:.1f} %)"),
    ]
    for label, value in summary_rows:
        cells = summary.add_row().cells
        cells[0].text, cells[1].text = label, str(value)

    document.add_heading("Betriebspunkt", level=2)
    operating_rows = [
        ("VDS", analysis_input["vds_v"], "V"), ("ID", analysis_input["id_a"], "A"),
        ("Pulsdauer", analysis_input["pulse_duration_s"] * 1e6, "µs"),
        ("Duty Cycle", analysis_input["duty_cycle"] * 100, "%"),
        ("Frequenz", analysis_input["frequency_hz"] / 1000, "kHz"),
        ("Gate-Spannung", analysis_input["gate_drive_voltage_v"], "V"),
        ("Eon", analysis_input["e_on_j"] * 1e6, "µJ"), ("Eoff", analysis_input["e_off_j"] * 1e6, "µJ"),
    ]
    operating = document.add_table(rows=1, cols=3)
    operating.style = "Light Grid Accent 1"
    operating.rows[0].cells[0].text, operating.rows[0].cells[1].text, operating.rows[0].cells[2].text = "Parameter", "Wert", "Einheit"
    for label, value, unit in operating_rows:
        cells = operating.add_row().cells
        cells[0].text, cells[1].text, cells[2].text = label, f"{value:.4g}", unit

    document.add_heading("Verlustaufteilung", level=2)
    for label, value in [("Leitverluste", result["p_conduction_w"]), ("Schaltverluste", result["p_switching_w"]), ("Gate-Drive-Verluste", result["p_gate_w"]), ("Gesamt", result["p_total_w"])]:
        document.add_paragraph(f"{label}: {value:.4g} W", style="List Bullet")

    document.add_heading("Thermischer Pfad und Grenzen", level=2)
    document.add_paragraph(
        f"Referenz: {analysis_input['temperature_reference']} bei {analysis_input['temperature_c']:.1f} °C; "
        f"Junction-Temperatur: {result['tj_c']:.1f} °C; Reserve: {result['temperature_margin_c']:.1f} K."
    )
    checks = document.add_table(rows=1, cols=2)
    checks.style = "Light Grid Accent 1"
    checks.rows[0].cells[0].text, checks.rows[0].cells[1].text = "Engineering Check", "Status"
    for label, passed in result.get("checks", {}).items():
        cells = checks.add_row().cells
        cells[0].text, cells[1].text = label.upper(), "PASS" if passed else "FAIL"

    source = payload.get("source", {})
    document.add_heading("Datenherkunft & Traceability", level=2)
    for label, value in [("Dataset", source.get("dataset_version")), ("C-Engine", source.get("engine_version")), ("Verifikation", source.get("verification_status")), ("Revision", source.get("revision"))]:
        document.add_paragraph(f"{label}: {value or '—'}", style="List Bullet")

    metadata = payload.get("analysis_metadata", {})
    for heading, key in [("Warnungen", "warnings"), ("Annahmen", "assumptions"), ("Modellgrenzen", "model_limitations")]:
        document.add_heading(heading, level=2)
        items = metadata.get(key, []) or (["Keine zusätzlichen Warnungen."] if key == "warnings" else ["Keine Angaben."])
        for item in items:
            document.add_paragraph(str(item), style="List Bullet")

    document.add_paragraph("Not certification software · Validate against datasheet and laboratory measurements")
    stream = io.BytesIO()
    document.save(stream)
    return stream.getvalue()


def export_record(record: dict[str, Any], format_name: str) -> Response:
    safe_name = "".join(character if character.isalnum() or character in "-_" else "_" for character in record.get("name", "analysis"))[:80]
    if format_name == "xlsx": data, media = workbook_for_analysis(record), "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    elif format_name == "pdf": data, media = pdf_for_analysis(record), "application/pdf"
    elif format_name == "docx": data, media = docx_for_analysis(record), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    elif format_name == "json": data, media = json.dumps(record, indent=2).encode(), "application/json"
    elif format_name == "csv":
        stream = io.StringIO(); writer = csv.writer(stream); writer.writerow(["field", "value", "unit"])
        payload = record["result"]; result = payload["result"]
        for row in [["analysis_id", record.get("id", ""), ""], ["name", record.get("name", ""), ""],
                    ["component", record["input"]["transistor_id"], ""], ["status", result["status"], ""],
                    ["vds", record["input"]["vds_v"], "V"], ["id", record["input"]["id_a"], "A"],
                    ["tj", result["tj_c"], "°C"], ["total_loss", result["p_total_w"], "W"],
                    ["closest_constraint", result["closest_constraint"]["type"], ""],
                    ["closest_reserve", result["closest_constraint"]["reserve_percent"], "%"]]: writer.writerow(row)
        data, media = stream.getvalue().encode("utf-8-sig"), "text/csv"
    else: raise HTTPException(status_code=422, detail="Supported formats: xlsx, pdf, docx, csv, json")
    return Response(content=data, media_type=media, headers={"Content-Disposition": f'attachment; filename="{safe_name}.{format_name}"'})


@app.get("/api/analyses/{analysis_id}/export/{format_name}")
def export_saved_analysis(analysis_id: str, format_name: str) -> Response:
    return export_record(get_analysis_or_404(analysis_id), format_name.lower())


@app.post("/api/export/{format_name}")
def export_current_analysis(format_name: str, request: SaveAnalysisRequest) -> Response:
    record = {"id": "unsaved", "name": request.name, "input": request.input.model_dump(), "result": request.result,
              "parent_id": request.parent_id, "report_config": request.report_config.model_dump() if request.report_config else None,
              "created_at": utc_now(), "updated_at": utc_now()}
    return export_record(record, format_name.lower())


@app.get("/api/batch/template")
def batch_template() -> Response:
    workbook = Workbook(); sheet = workbook.active; sheet.title = "INPUT"
    headers = ["Analysis Name", "Component / OPN", "Mode", "VDS [V]", "ID [A]", "Pulse Duration [µs]", "Duty Cycle [%]",
               "Frequency [kHz]", "Gate Voltage [V]", "Eon [µJ]", "Eoff [µJ]", "Case Temperature [°C]",
               "Ambient Temperature [°C]", "Thermal Mode", "RthCS [K/W]", "RthSA [K/W]", "Safety Factor"]
    sheet.append(headers); sheet.append(["Example", "CSD19536KTT", "SWITCHING", 48, 40, 10, 50, 100, 10, 20, 15, 25, "", "CASE", 0, 0, 1.2])
    sheet.freeze_panes="A2"; sheet.auto_filter.ref=sheet.dimensions
    for cell in sheet[1]: cell.font=Font(bold=True,color="FFFFFF"); cell.fill=PatternFill("solid",fgColor="0B5CAD")
    for column in range(1, len(headers)+1): sheet.column_dimensions[chr(64+column) if column <= 26 else "Q"].width=22
    instructions = workbook.create_sheet("INSTRUCTIONS")
    instructions.append(["Field", "Requirement", "Description"])
    for row in [["Component / OPN","Required","Must exist in the current TransiSafe dataset"], ["Mode","Required","LINEAR or SWITCHING"],
                ["VDS / ID / Pulse Duration / Duty Cycle","Required","Explicit units are shown in the INPUT header"],
                ["Frequency / Gate Voltage / Eon / Eoff","Switching","Required only in SWITCHING mode"],
                ["Case or Ambient Temperature","Mode-dependent","Provide the value matching Thermal Mode"],
                ["RthCS / RthSA","Ambient","Optional external thermal path values"], ["Safety Factor","Required","Must be at least 1.0"]]: instructions.append(row)
    stream=io.BytesIO(); workbook.save(stream)
    return Response(stream.getvalue(), media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    headers={"Content-Disposition": 'attachment; filename="TransiSafe-Batch-Template.xlsx"'})


FIELD_ALIASES = {
    "analysis_name": {"analysis name", "analyse", "name"}, "transistor_id": {"component / opn", "component", "opn", "mosfet", "modell"},
    "mode": {"mode", "modus"}, "vds_v": {"vds [v]", "vds", "drain voltage"}, "id_a": {"id [a]", "id", "drain current"},
    "pulse_duration_us": {"pulse duration [µs]", "pulse duration", "pulsdauer"}, "duty_cycle_percent": {"duty cycle [%]", "duty cycle"},
    "frequency_khz": {"frequency [khz]", "frequency", "switch frequency"}, "gate_drive_voltage_v": {"gate voltage [v]", "gate voltage", "vgs"},
    "e_on_uj": {"eon [µj]", "eon"}, "e_off_uj": {"eoff [µj]", "eoff"}, "case_temperature_c": {"case temperature [°c]", "case temp", "tc"},
    "ambient_temperature_c": {"ambient temperature [°c]", "ambient temp", "ta"}, "temperature_reference": {"thermal mode", "temperature reference"},
    "rth_cs_k_per_w": {"rthcs [k/w]", "rthcs"}, "rth_sa_k_per_w": {"rthsa [k/w]", "rthsa"}, "safety_factor": {"safety factor", "sicherheitsfaktor"},
}


def suggested_mapping(columns: list[str]) -> dict[str, str | None]:
    mapping: dict[str, str | None] = {}
    for column in columns:
        normalized = column.strip().lower(); matches = [field for field, aliases in FIELD_ALIASES.items() if normalized in aliases]
        mapping[column] = matches[0] if len(matches) == 1 else None
    return mapping


@app.post("/api/batch/preview")
def batch_preview(content: bytes = Body(media_type="application/octet-stream"), filename: str = Query(..., max_length=160)) -> dict[str, Any]:
    try:
        if filename.lower().endswith(".xlsx"):
            workbook=load_workbook(io.BytesIO(content), data_only=True, read_only=True); sheet=workbook[workbook.sheetnames[0]]
            values=list(sheet.iter_rows(values_only=True)); columns=[str(value or "") for value in values[0]]; rows=[list(row) for row in values[1:] if any(value is not None for value in row)]
        elif filename.lower().endswith(".csv"):
            decoded=content.decode("utf-8-sig"); sample=decoded[:4096]; dialect=csv.Sniffer().sniff(sample, delimiters=",;\t")
            reader=csv.reader(io.StringIO(decoded), dialect); values=list(reader); columns=values[0]; rows=values[1:]
        else: raise HTTPException(status_code=422, detail="Only .xlsx and .csv files are supported")
    except HTTPException: raise
    except Exception as exc: raise HTTPException(status_code=422, detail=f"Unable to read batch file: {exc}") from exc
    mapping=suggested_mapping(columns)
    return {"filename": filename, "columns": columns, "suggested_mapping": mapping,
            "row_count": len(rows), "preview_rows": [dict(zip(columns,row)) for row in rows[:20]], "raw_rows": [dict(zip(columns,row)) for row in rows]}


def normalized_batch_input(raw: dict[str, Any], mapping: dict[str, str | None]) -> tuple[str, AnalysisRequest]:
    values={target: raw.get(source) for source,target in mapping.items() if target}
    mode=str(values.get("mode","")).strip().upper(); thermal=str(values.get("temperature_reference", "")).strip().upper()
    if thermal not in {"CASE","AMBIENT"}: raise ValueError("Missing or unsupported thermal mode; use CASE or AMBIENT")
    temperature = values.get("case_temperature_c") if thermal == "CASE" else values.get("ambient_temperature_c")
    def number(name: str, default: float | None=None) -> float:
        value=values.get(name, default)
        if value is None or str(value).strip()=="":
            if default is None: raise ValueError(f"Missing required value: {name}")
            return default
        return float(str(value).replace(",","."))
    request=AnalysisRequest(transistor_id=str(values.get("transistor_id","")).strip(), mode=mode,
        vds_v=number("vds_v"), id_a=number("id_a"), pulse_duration_s=number("pulse_duration_us")/1e6,
        duty_cycle=number("duty_cycle_percent")/100, frequency_hz=number("frequency_khz",0)*1000,
        gate_drive_voltage_v=number("gate_drive_voltage_v",0), e_on_j=number("e_on_uj",0)/1e6, e_off_j=number("e_off_uj",0)/1e6,
        temperature_reference=thermal, temperature_c=float(str(temperature).replace(",",".")) if temperature not in (None,"") else number("case_temperature_c" if thermal=="CASE" else "ambient_temperature_c"),
        rth_cs_k_per_w=number("rth_cs_k_per_w",0), rth_sa_k_per_w=number("rth_sa_k_per_w",0), safety_factor=number("safety_factor"))
    return str(values.get("analysis_name","")).strip(), request


@app.post("/api/batch/validate")
def validate_batch(payload: dict[str, Any]) -> dict[str, Any]:
    rows=payload.get("rows",[]); mapping=payload.get("mapping",{}); valid=[]; invalid=[]; known={m["id"] for m in load_models()}
    for index, raw in enumerate(rows,2):
        try:
            name, request=normalized_batch_input(raw,mapping)
            if request.transistor_id not in known: raise ValueError(f"Unknown component: {request.transistor_id}")
            valid.append({"source_row":index,"analysis_name":name or f"Row {index}","input":request.model_dump()})
        except Exception as exc: invalid.append({"source_row":index,"errors":[str(exc)]})
    return {"detected":len(rows),"ready":len(valid),"warnings":0,"invalid":len(invalid),"valid_rows":valid,"invalid_rows":invalid}


@app.post("/api/batch/analyze")
def analyze_batch(request: BatchAnalyzeRequest) -> dict[str, Any]:
    results=[]
    for row in request.rows:
        try: results.append({"source_row":row.source_row,"analysis_name":row.analysis_name,"input":row.input.model_dump(),"result":run_engine(row.input)})
        except HTTPException as exc: results.append({"source_row":row.source_row,"analysis_name":row.analysis_name,"input":row.input.model_dump(),"error":str(exc.detail)})
    return {"count":len(results),"results":results}


@app.post("/api/batch/export")
def export_batch(payload: dict[str, Any]) -> Response:
    rows = payload.get("results", [])
    if not isinstance(rows, list) or not rows: raise HTTPException(status_code=422, detail="No batch results supplied")
    workbook=Workbook(); workbook.remove(workbook.active)
    result_rows=[["Source Row","Analysis","Component","Mode","VDS [V]","ID [A]","Temperature [°C]","Assessment","Relevant Limit","Reserve [%]","SOA Limit [A]","Tj [°C]","Thermal Reserve [K]","Total Loss [W]","Voltage Reserve [%]","Current Reserve [%]","Warnings"]]
    input_rows=[["Source Row","Analysis","Component","Mode","VDS [V]","ID [A]","Pulse Duration [s]","Duty Cycle","Frequency [Hz]","Gate Voltage [V]","Eon [J]","Eoff [J]","Temperature Reference","Temperature [°C]","RthCS [K/W]","RthSA [K/W]","Safety Factor"]]
    warning_rows=[["Source Row","Analysis","Warning"]]; trace_rows=[["Source Row","Component","Datasheet Revision","Retrieved Date","Dataset Version","Verification Status","Engine Version"]]
    for row in rows:
        inp=row.get("input",{}); analysis=row.get("result"); source_row=row.get("source_row"); name=row.get("analysis_name","")
        input_rows.append([source_row,name,inp.get("transistor_id"),inp.get("mode"),inp.get("vds_v"),inp.get("id_a"),inp.get("pulse_duration_s"),inp.get("duty_cycle"),inp.get("frequency_hz"),inp.get("gate_drive_voltage_v"),inp.get("e_on_j"),inp.get("e_off_j"),inp.get("temperature_reference"),inp.get("temperature_c"),inp.get("rth_cs_k_per_w"),inp.get("rth_sa_k_per_w"),inp.get("safety_factor")])
        if not analysis:
            result_rows.append([source_row,name,inp.get("transistor_id"),inp.get("mode"),inp.get("vds_v"),inp.get("id_a"),inp.get("temperature_c"),"ERROR","","","","","","","","",row.get("error","")]); warning_rows.append([source_row,name,row.get("error","")]); continue
        result=analysis["result"]; source=analysis["source"]; warnings=analysis.get("analysis_metadata",{}).get("warnings",[])
        result_rows.append([source_row,name,inp.get("transistor_id"),inp.get("mode"),inp.get("vds_v"),inp.get("id_a"),inp.get("temperature_c"),result["status"],result["closest_constraint"]["type"],result["closest_constraint"]["reserve_percent"],result["soa_limit_a"],result["tj_c"],result["temperature_margin_c"],result["p_total_w"],result["margins"]["voltage_reserve_percent"],result["margins"]["current_reserve_percent"]," | ".join(warnings)])
        for warning in warnings: warning_rows.append([source_row,name,warning])
        trace_rows.append([source_row,inp.get("transistor_id"),source["revision"],source["retrieved_date"],source["dataset_version"],source["verification_status"],source["engine_version"]])
    statuses={str(row[7]) for row in result_rows[1:]}; summary_rows=[["Metric","Value"],["Operating Points",len(rows)]]+[[status,sum(1 for row in result_rows[1:] if row[7]==status)] for status in sorted(statuses)]
    add_sheet(workbook,"SUMMARY",summary_rows); results_sheet=add_sheet(workbook,"RESULTS",result_rows,(12,28,24,15,13,13,18,20,22,16,16,14,18,16,18,18,55)); add_sheet(workbook,"INPUTS",input_rows,(12,28,24,15,13,13,20,15,18,18,16,16,20,18,16,16,14)); add_sheet(workbook,"WARNINGS",warning_rows,(12,30,90)); add_sheet(workbook,"TRACEABILITY",trace_rows,(12,24,30,18,18,24,18)); add_sheet(workbook,"METADATA",[["Field","Value"],["Exported",utc_now()],["Application Version",APP_VERSION],["Dataset Version",DATASET_VERSION],["Source","Native C engine results"]])
    for row in range(2, results_sheet.max_row+1): results_sheet.cell(row,8).fill=status_fill(str(results_sheet.cell(row,8).value))
    stream=io.BytesIO();workbook.save(stream)
    return Response(stream.getvalue(),media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",headers={"Content-Disposition":'attachment; filename="TransiSafe-Batch-Results.xlsx"'})
